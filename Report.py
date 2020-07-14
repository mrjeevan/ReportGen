import docx
import csv
import os
from datetime import datetime
from docx.shared import Inches

print('Ver 1.0')

doc = docx.Document('Template.docx')
t = doc.tables[0]
pre2 = {}
save = datetime.now().date()

default = {}
with open('list.txt','r') as f:
	w = f.readlines()

	for row in w:
		a = row.split(',')
		default[a[0]] = a[1].replace('\n','')



	for x,row in enumerate(t.rows):
		for y,cell in enumerate(row.cells):
			s = cell.text
			# print(f'***debug : {s}***')
			if s == 'I/P Date:':			 	# AutoFill Date
				cell.text = str(save)
				pre2[s] = str(save)
			elif s not in pre2.keys() and  'I/P' in s:	# User Input or AutoFill
				ans = input(cell.text + ' ')
				rrr = cell.text
				if ans != '':
					cell.text = ans
					print('\t\t Please Change it in list.txt also.... ')
				else:							# Auto Fill From File
					print(f'***Set To Default! : ',end = '')
					cell.text = default[s]			
					print(cell.text)
			elif 'Problem' in s  and 'Problem' not in pre2.keys(): # 'Problem Statement :'
				ans = input('Problem Statement : ')
				cell.text = 'Problem Statement : ' + str(ans)
				pre2['Problem'] = True
			elif s == 'Status:':
				ans = input('Status : ')
				cell.text = 'Status : ' + str(ans)
				pre2[cell.text] = [x,y]
				# print(f' Debug {cell.text}[{x},{y}]')


## Adding Screen-Shots .........



# print(f' *** debug ***pwd : {os.getcwd()}')
os.chdir('ScreenShots')
folder = os.listdir(os.getcwd())
# print(f'***debug*** folder contains {folder}' )
if len(folder) == 0:
	print('***************Add pic to ScreenShots folder !*****************')
else :
	for pic in folder:
		# prepare path
		PATH = os.path.join(os.getcwd(),pic)                        
		# Add pic
		doc.add_picture(PATH, width=Inches(6.7), height=Inches(4.0))
		# add description !
		description = input(f'describe {pic} : ')
		doc.add_paragraph(description)

os.chdir('..')


doc.save(str(save)+'.docx')     # Autogenerate docx file




# EOF
